# app.py
# Birchwood Credit Report Analyzer (Streamlit)
# - Upload Birchwood PDF
# - Extract negative accounts (collections/charge-offs/serious lates)
# - Generate ranked repair plan + Word doc

import re
import io
import tempfile
from dataclasses import dataclass, asdict
from typing import List, Optional, Tuple, Dict

import streamlit as st

# PDF extraction
import fitz  # PyMuPDF

# OCR fallback
from PIL import Image
import pytesseract

# Word doc output
from docx import Document


# ----------------------------
# Models
# ----------------------------
@dataclass
class NegativeAccount:
    creditor: str
    account_number: str
    acct_type: str  # "COLLECTION" | "CHARGE OFF" | "SERIOUS DELINQUENCY"
    balance: Optional[str] = None
    past_due: Optional[str] = None
    last_reported: Optional[str] = None
    opened: Optional[str] = None
    dla: Optional[str] = None
    original_creditor: Optional[str] = None
    source_bureaus: Optional[str] = None
    notes: Optional[str] = None
    contact_phone: Optional[str] = None
    contact_address: Optional[str] = None
    strategy: Optional[str] = None
    priority_rank: Optional[int] = None


# ----------------------------
# Utilities
# ----------------------------
def normalize_spaces(s: str) -> str:
    return re.sub(r"[ \t]+", " ", s).strip()

def money_to_float(m: str) -> Optional[float]:
    if not m:
        return None
    m = m.replace(",", "").replace("$", "").strip()
    try:
        return float(m)
    except:
        return None

def safe_text(s: Optional[str]) -> str:
    return s if s else ""

def looks_like_birchwood(text: str) -> bool:
    t = text.lower()
    return ("birchwood" in t and "merged infile credit report" in t) or ("birchwood credit services" in t)

def extract_text_pymupdf(pdf_bytes: bytes) -> Tuple[str, List[str]]:
    """Return full text and per-page text."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages = []
    for i in range(doc.page_count):
        pages.append(doc.load_page(i).get_text("text") or "")
    full = "\n\n".join(pages)
    return full, pages

def ocr_pdf_pages(pdf_bytes: bytes, max_pages: int = 12, dpi: int = 200) -> Tuple[str, List[str]]:
    """OCR pages to text (fallback for image-based PDFs)."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    n = min(doc.page_count, max_pages)
    page_texts: List[str] = []
    for i in range(n):
        page = doc.load_page(i)
        mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        txt = pytesseract.image_to_string(img)
        page_texts.append(txt or "")
    full = "\n\n".join(page_texts)
    return full, page_texts

def choose_best_text(primary_full: str, primary_pages: List[str],
                     ocr_full: str, ocr_pages: List[str]) -> Tuple[str, List[str], str]:
    """Pick the best extraction result based on heuristics."""
    # Score: length + presence of key tokens
    def score(t: str) -> int:
        tl = t.lower()
        s = len(t)
        for k in ["tradelines", "collections", "charge off", "derogatory summary", "creditors"]:
            if k in tl:
                s += 2000
        if looks_like_birchwood(t):
            s += 5000
        return s

    p_score = score(primary_full)
    o_score = score(ocr_full)

    if o_score > p_score * 1.1 and len(ocr_full.strip()) > 500:
        return ocr_full, ocr_pages, "OCR"
    return primary_full, primary_pages, "TEXT"


# ----------------------------
# Parsing Birchwood structures
# ----------------------------
def parse_creditors_directory(full_text: str) -> Dict[str, Dict[str, str]]:
    """
    Parse the CREDITORS directory section:
    Lines often like:
      "PORTFOLIO RECOVERY A 120 CORPORATE BLVD... 888-772-7326"
    We'll map by a normalized creditor key to phone/address when we can.
    """
    directory: Dict[str, Dict[str, str]] = {}
    t = full_text

    # Try to isolate CREDITORS section
    m = re.search(r"\nCREDITORS\s*\n(.+?)(\n{2,}|\Z)", t, flags=re.IGNORECASE | re.DOTALL)
    if not m:
        # Some reports keep it across pages; be more permissive
        m = re.search(r"CREDITORS\s*\n(.+)", t, flags=re.IGNORECASE | re.DOTALL)
    if not m:
        return directory

    blob = m.group(1)
    lines = [normalize_spaces(x) for x in blob.splitlines() if normalize_spaces(x)]
    # Stop if we hit something clearly not directory
    stop_tokens = {"miscellaneous", "end of report", "www.", "experian", "transunion", "equifax"}
    cleaned = []
    for line in lines:
        low = line.lower()
        if any(tok in low for tok in stop_tokens):
            break
        cleaned.append(line)

    # Heuristic parse: creditor name then address then phone at end
    for line in cleaned:
        # Look for phone at end
        phone = None
        phone_m = re.search(r"(\(?\d{3}\)?[-.\s]\d{3}[-.\s]\d{4}|\d{3}[-.\s]\d{3}[-.\s]\d{4})\s*$", line)
        if phone_m:
            phone = phone_m.group(1).replace(".", "-").replace(" ", "-")
            pre = line[:phone_m.start()].strip()
        else:
            pre = line.strip()

        # Split pre into name + address guess
        # Name is leading chunk until we see a digit (street number) or PO BOX
        name = pre
        addr = ""
        street_m = re.search(r"\b(PO BOX|P O BOX|BOX|\d{1,6}\s)\b", pre, flags=re.IGNORECASE)
        if street_m:
            name = pre[:street_m.start()].strip(" -")
            addr = pre[street_m.start():].strip()
        key = re.sub(r"[^a-z0-9]+", "", name.lower())

        if len(key) >= 4:
            directory[key] = {
                "name": name.strip(),
                "address": addr.strip(),
                "phone": phone or ""
            }
    return directory

def parse_negative_accounts(full_text: str) -> List[NegativeAccount]:
    """
    Parse TRADELINES blocks for:
    - COLLECTION
    - CHARGE OFF
    Also capture creditor name & account number line preceding, plus key fields.
    """
    text = full_text

    # Extract trade lines chunk(s) after "TRADELINES"
    # We'll scan globally for patterns
    accounts: List[NegativeAccount] = []

    # Pattern: creditor name line then account number line then fields (Opened/Reported/Balance/etc)
    # We'll use a rolling window approach for each match of COLLECTION/CHARGE OFF.
    tokens = [m.start() for m in re.finditer(r"\b(COLLECTION|CHARGE OFF)\b", text, flags=re.IGNORECASE)]
    for pos in tokens:
        window_start = max(0, pos - 500)
        window_end = min(len(text), pos + 900)
        win = text[window_start:window_end]

        acct_type_m = re.search(r"\b(COLLECTION|CHARGE OFF)\b", win, flags=re.IGNORECASE)
        if not acct_type_m:
            continue
        acct_type = acct_type_m.group(1).upper()

        # Try to get creditor + account number above type
        # Grab up to 10 lines before the type
        pre = win[:acct_type_m.start()]
        pre_lines = [normalize_spaces(x) for x in pre.splitlines() if normalize_spaces(x)]
        creditor = ""
        acctno = ""

        # Heuristic: creditor is one of the last 6 lines, account number is the next line with lots of digits
        for i in range(len(pre_lines) - 1, max(-1, len(pre_lines) - 10), -1):
            line = pre_lines[i]
            if not acctno:
                # Account numbers often have 8+ digits or alnum ids
                if re.search(r"\b[0-9A-Z]{8,}\b", line) and not re.search(r"\b(Opened|Reported|Balance|Payment|ECOA|Source)\b", line, re.IGNORECASE):
                    acctno = re.search(r"\b[0-9A-Z]{8,}\b", line).group(0)
                    continue
            if not creditor:
                if not re.search(r"\b(Opened|Reported|Hi\.|Credit|Limit|Reviewed|DLA|ECOA|Source|30-59|60-89|90\+|Past Due|Payment|Balance)\b", line, re.IGNORECASE):
                    # creditor names usually have letters
                    if re.search(r"[A-Za-z]", line):
                        creditor = line
                        break

        # Fields in window
        opened = find_field(win, r"\bOpened\s*([0-9]{2}/[0-9]{2})")
        reported = find_field(win, r"\bReported\s*([0-9]{2}/[0-9]{2})")
        balance = find_money_field(win, r"\bBalance\s*\$?([0-9,]+)")
        past_due = find_money_field(win, r"\bPast Due\s*\$?([0-9,]+)")
        dla = find_field(win, r"\bDLA\s*([0-9]{2}/[0-9]{2})")
        orig = find_field(win, r"ORIGINAL CREDITOR:\s*([A-Z0-9&\.\-\s]+)")
        source_bureaus = find_field(win, r"\bSource\s*\(B\)\s*([A-Z/]+)")

        notes = []
        # Common flags
        for note_pat in [
            r"CONSUMER DISPUTES THIS ACCOUNT INFORMATION",
            r"FACTORING COMPANY;.*",
            r"PROFIT AND LOSS WRITEOFF;.*",
            r"TRANSFER/SOLD.*",
            r"CHARGED OFF ACCOUNT.*",
            r"COLLECTION ACCOUNT.*",
        ]:
            nm = re.search(note_pat, win, flags=re.IGNORECASE)
            if nm:
                notes.append(normalize_spaces(nm.group(0)))

        # Avoid duplicates by (creditor, acctno, type, balance)
        key = (creditor.strip().lower(), acctno.strip().lower(), acct_type, safe_text(balance))
        if any((a.creditor.strip().lower(), a.account_number.strip().lower(), a.acct_type, safe_text(a.balance)) == key for a in accounts):
            continue

        accounts.append(NegativeAccount(
            creditor=creditor or "UNKNOWN",
            account_number=acctno or "UNKNOWN",
            acct_type=acct_type,
            balance=balance,
            past_due=past_due,
            last_reported=reported,
            opened=opened,
            dla=dla,
            original_creditor=orig,
            source_bureaus=source_bureaus,
            notes="; ".join(notes) if notes else None,
        ))

    return accounts

def find_field(text: str, pattern: str) -> Optional[str]:
    m = re.search(pattern, text, flags=re.IGNORECASE)
    if not m:
        return None
    return normalize_spaces(m.group(1))

def find_money_field(text: str, pattern: str) -> Optional[str]:
    m = re.search(pattern, text, flags=re.IGNORECASE)
    if not m:
        return None
    return "$" + m.group(1).replace(",", "").strip().replace(" ", "")

def parse_scores(full_text: str) -> Dict[str, Optional[int]]:
    scores = {"TransUnion": None, "Experian": None, "Equifax": None}
    # Try several patterns
    for bureau in ["TRANSUNION", "EXPERIAN", "EQUIFAX"]:
        m = re.search(rf"{bureau}.*?SCORE:\s*([0-9]{{3}})", full_text, flags=re.IGNORECASE | re.DOTALL)
        if m:
            val = int(m.group(1))
            if bureau == "TRANSUNION": scores["TransUnion"] = val
            if bureau == "EXPERIAN": scores["Experian"] = val
            if bureau == "EQUIFAX": scores["Equifax"] = val
    return scores

def rank_and_attach_strategies(accounts: List[NegativeAccount], directory: Dict[str, Dict[str, str]]) -> List[NegativeAccount]:
    """
    Rank:
      1) Collections (esp smaller balances) first
      2) Charge-offs next (smaller balances first)
    Attach:
      - contact phone/address from directory if match
      - strategy text
    """
    def dir_lookup(name: str) -> Dict[str, str]:
        key = re.sub(r"[^a-z0-9]+", "", name.lower())
        # try exact, then partial
        if key in directory:
            return directory[key]
        for k, v in directory.items():
            if key and (key in k or k in key):
                return v
        return {}

    # Compute sort key
    def sort_key(a: NegativeAccount):
        bal = money_to_float(a.balance) if a.balance else 999999.0
        if a.acct_type == "COLLECTION":
            return (0, bal)
        if a.acct_type == "CHARGE OFF":
            return (1, bal)
        return (2, bal)

    sorted_acc = sorted(accounts, key=sort_key)

    for idx, a in enumerate(sorted_acc, start=1):
        a.priority_rank = idx
        d = dir_lookup(a.creditor)
        if d:
            a.contact_phone = d.get("phone") or None
            a.contact_address = d.get("address") or None

        if a.acct_type == "COLLECTION":
            # Offer ~35-60% settlement if not tiny; tiny often can do 60-80% for fast delete
            bal = money_to_float(a.balance) if a.balance else None
            if bal is not None and bal <= 300:
                settle = "60–80%"
            else:
                settle = "35–55%"
            a.strategy = (
                f"Call and negotiate a Pay-For-Delete. Target settlement around {settle} of balance. "
                f"Ask for deletion from all three bureaus in writing (email/letter) before paying."
            )
        elif a.acct_type == "CHARGE OFF":
            a.strategy = (
                "Charge-offs rarely delete on request. Best play: negotiate settlement (or pay in full if needed for underwriting), "
                "then follow with a goodwill request OR dispute any inaccurate payment history/balance/status. "
                "If a third-party collector also reports, prioritize deleting the collection tradeline."
            )
        else:
            a.strategy = "Investigate and address serious delinquency with goodwill/dispute as appropriate."

    return sorted_acc

def build_word_doc(client_name: str, scores: Dict[str, Optional[int]],
                   accounts: List[NegativeAccount], extraction_mode: str) -> bytes:
    doc = Document()
    doc.add_heading(f"Credit Repair Action Plan – {client_name}", level=1)

    doc.add_paragraph(f"Extraction mode used: {extraction_mode}")
    doc.add_heading("Current Scores", level=2)
    doc.add_paragraph(
        f"TransUnion: {scores.get('TransUnion') or 'N/A'}\n"
        f"Experian: {scores.get('Experian') or 'N/A'}\n"
        f"Equifax: {scores.get('Equifax') or 'N/A'}"
    )

    doc.add_heading("Priority Repair List (Easiest → Hardest)", level=2)

    for a in accounts:
        title = f"{a.priority_rank}. {a.creditor} — {a.acct_type}"
        doc.add_heading(title, level=3)

        doc.add_paragraph(f"Account Number: {a.account_number}")
        if a.original_creditor:
            doc.add_paragraph(f"Original Creditor: {a.original_creditor}")
        if a.balance:
            doc.add_paragraph(f"Balance: {a.balance}")
        if a.past_due and a.past_due != "$0":
            doc.add_paragraph(f"Past Due: {a.past_due}")
        if a.opened:
            doc.add_paragraph(f"Opened: {a.opened}")
        if a.last_reported:
            doc.add_paragraph(f"Last Reported: {a.last_reported}")
        if a.dla:
            doc.add_paragraph(f"DLA: {a.dla}")
        if a.source_bureaus:
            doc.add_paragraph(f"Reporting (Source): {a.source_bureaus}")

        if a.contact_phone or a.contact_address:
            doc.add_paragraph(f"Contact Phone: {a.contact_phone or 'N/A'}")
            doc.add_paragraph(f"Contact Address: {a.contact_address or 'N/A'}")

        if a.notes:
            doc.add_paragraph(f"Notes: {a.notes}")

        doc.add_paragraph(f"Strategy: {a.strategy or ''}")

    doc.add_heading("30–60 Day Timeline", level=2)
    doc.add_paragraph(
        "Week 1: Resolve the smallest collections first (highest deletion probability).\n"
        "Week 2: Negotiate remaining collections; confirm deletions in writing.\n"
        "Week 3: Address charge-offs (settlement / pay-in-full if needed) + send goodwill letters.\n"
        "Week 4: Dispute any inaccurate items; monitor for deletions and updates."
    )

    doc.add_heading("Score Boost Tips", level=2)
    doc.add_paragraph(
        "- Keep revolving utilization under 10%.\n"
        "- Avoid new inquiries unless needed.\n"
        "- Make every payment on time.\n"
        "- Consider adding a strong authorized user tradeline (seasoned card) if appropriate."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ----------------------------
# Streamlit App
# ----------------------------
st.set_page_config(page_title="Birchwood Credit Report Analyzer", layout="wide")
st.title("Birchwood Credit Report Analyzer")
st.caption("Upload a Birchwood 'Merged Infile Credit Report' PDF → get negative accounts + ranked repair plan + Word doc.")

colA, colB = st.columns([2, 1])
with colB:
    st.subheader("Settings")
    use_ocr = st.checkbox("Force OCR (slower, for scanned PDFs)", value=False)
    ocr_max_pages = st.slider("OCR max pages", 1, 25, 12)
    st.write("Tip: leave OCR off unless the report is image-only.")

with colA:
    uploaded = st.file_uploader("Upload Birchwood credit report PDF", type=["pdf"])

if not uploaded:
    st.stop()

pdf_bytes = uploaded.read()

# Extract
primary_full, primary_pages = extract_text_pymupdf(pdf_bytes)
ocr_full, ocr_pages = ("", [])
mode = "TEXT"
if use_ocr or len(primary_full.strip()) < 500:
    ocr_full, ocr_pages = ocr_pdf_pages(pdf_bytes, max_pages=ocr_max_pages)

full_text, page_texts, extraction_mode = choose_best_text(primary_full, primary_pages, ocr_full, ocr_pages)

if not looks_like_birchwood(full_text) and not use_ocr:
    st.warning("This doesn't look like a Birchwood merged infile report, but I’ll still try to parse it.")

# Parse basics
scores = parse_scores(full_text)
directory = parse_creditors_directory(full_text)

raw_accounts = parse_negative_accounts(full_text)
ranked = rank_and_attach_strategies(raw_accounts, directory)

# Client name (best effort)
client_name = "Client"
name_m = re.search(r"\bAPPLICANT\s+([A-Z][A-Z'\-]+,\s*[A-Z][A-Z'\-]+)", full_text, flags=re.IGNORECASE)
if name_m:
    client_name = normalize_spaces(name_m.group(1).title())

st.subheader("Summary")
c1, c2, c3, c4 = st.columns(4)
c1.metric("TransUnion", scores.get("TransUnion") or 0)
c2.metric("Experian", scores.get("Experian") or 0)
c3.metric("Equifax", scores.get("Equifax") or 0)
c4.metric("Negative Accounts Found", len(ranked))

st.caption(f"Extraction used: **{extraction_mode}**")

if len(ranked) == 0:
    st.error("No collections or charge-offs detected. If this is a scanned report, try enabling OCR.")
    st.stop()

st.subheader("Priority Repair List (Easiest → Hardest)")
for a in ranked:
    with st.expander(f"{a.priority_rank}. {a.creditor} — {a.acct_type} — Balance: {a.balance or 'N/A'}", expanded=False):
        st.write({
            "Creditor": a.creditor,
            "Account Number": a.account_number,
            "Type": a.acct_type,
            "Balance": a.balance,
            "Past Due": a.past_due,
            "Opened": a.opened,
            "Last Reported": a.last_reported,
            "DLA": a.dla,
            "Original Creditor": a.original_creditor,
            "Reporting": a.source_bureaus,
            "Contact Phone": a.contact_phone,
            "Contact Address": a.contact_address,
            "Notes": a.notes,
        })
        st.markdown(f"**Strategy:** {a.strategy}")

st.subheader("Download Word Doc Action Plan")
doc_bytes = build_word_doc(client_name=client_name, scores=scores, accounts=ranked, extraction_mode=extraction_mode)

st.download_button(
    label="Download Credit Repair Action Plan (.docx)",
    data=doc_bytes,
    file_name=f"{re.sub(r'[^a-zA-Z0-9]+','_',client_name).strip('_')}_credit_repair_action_plan.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

with st.expander("Debug (extracted text preview)"):
    st.text(full_text[:6000])
